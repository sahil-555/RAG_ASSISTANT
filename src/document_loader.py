"""
Document loader for the RAG Assistant project.

This module handles loading documents from various formats and preparing them
for ingestion into the vector store.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass

import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata for a document."""
    source: str
    file_type: str
    file_size: int
    page_count: Optional[int] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None


class DocumentLoader:
    """Handles loading and processing documents from various formats."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document loader.
        
        Args:
            chunk_size: Size of text chunks for splitting
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Supported file extensions
        self.supported_extensions = {
            '.txt': self._load_text_file,
            '.pdf': self._load_pdf_file,
            '.docx': self._load_docx_file,
            '.md': self._load_text_file,
        }
    
    def load_documents(self, directory_path: str) -> List[LangChainDocument]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory_path: Path to the directory containing documents
            
        Returns:
            List of LangChain Document objects
        """
        directory = Path(directory_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        documents = []
        logger.info(f"Loading documents from: {directory_path}")
        
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    file_docs = self._load_single_file(file_path)
                    documents.extend(file_docs)
                    logger.info(f"Loaded {len(file_docs)} chunks from {file_path.name}")
                except Exception as e:
                    logger.error(f"Error loading {file_path}: {e}")
        
        logger.info(f"Total documents loaded: {len(documents)}")
        return documents
    
    def load_single_document(self, file_path: str) -> List[LangChainDocument]:
        """
        Load a single document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of LangChain Document objects (chunks)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        return self._load_single_file(file_path)
    
    def _load_single_file(self, file_path: Path) -> List[LangChainDocument]:
        """Load a single file and return its chunks."""
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Load the file content
        loader_func = self.supported_extensions[file_extension]
        content = loader_func(file_path)
        
        # Get metadata
        metadata = self._extract_metadata(file_path)
        
        # Split content into chunks
        chunks = self.text_splitter.split_text(content)
        
        # Create LangChain Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc = LangChainDocument(
                page_content=chunk,
                metadata={
                    **metadata,
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)
        
        return documents
    
    def _load_text_file(self, file_path: Path) -> str:
        """Load content from a text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _load_pdf_file(self, file_path: Path) -> str:
        """Load content from a PDF file."""
        content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            raise
        
        return content
    
    def _load_docx_file(self, file_path: Path) -> str:
        """Load content from a DOCX file."""
        try:
            doc = Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            return content
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            raise
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from a file."""
        try:
            stat = file_path.stat()
            return {
                "source": str(file_path),
                "file_type": file_path.suffix.lower(),
                "file_size": stat.st_size,
                "created_date": str(stat.st_ctime),
                "modified_date": str(stat.st_mtime),
            }
        except Exception as e:
            logger.warning(f"Could not extract metadata for {file_path}: {e}")
            return {
                "source": str(file_path),
                "file_type": file_path.suffix.lower(),
                "file_size": 0,
            }
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_extensions.keys())
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validate if a file can be processed.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if file can be processed, False otherwise
        """
        file_path = Path(file_path)
        return (
            file_path.exists() and 
            file_path.is_file() and 
            file_path.suffix.lower() in self.supported_extensions
        )


def create_sample_documents() -> List[LangChainDocument]:
    """Create sample documents for testing purposes."""
    sample_texts = [
        """
        Artificial Intelligence (AI) is a branch of computer science that aims to create 
        intelligent machines that work and react like humans. Some of the activities 
        computers with artificial intelligence are designed for include speech recognition, 
        learning, planning, and problem solving.
        """,
        """
        Machine Learning is a subset of AI that provides systems the ability to automatically 
        learn and improve from experience without being explicitly programmed. Machine learning 
        focuses on the development of computer programs that can access data and use it to 
        learn for themselves.
        """,
        """
        Deep Learning is a subset of machine learning that uses neural networks with multiple 
        layers to model and understand complex patterns. It's particularly effective for 
        tasks like image recognition, natural language processing, and speech recognition.
        """,
        """
        Natural Language Processing (NLP) is a field of AI that gives machines the ability 
        to read, understand, and derive meaning from human languages. It combines computational 
        linguistics with statistical, machine learning, and deep learning models.
        """,
        """
        Vector databases are specialized databases designed to store, manage, and search 
        high-dimensional vector data efficiently. They're essential for AI applications 
        that need to perform similarity searches, such as recommendation systems and 
        semantic search engines.
        """
    ]
    
    documents = []
    for i, text in enumerate(sample_texts):
        doc = LangChainDocument(
            page_content=text.strip(),
            metadata={
                "source": f"sample_doc_{i+1}.txt",
                "file_type": ".txt",
                "file_size": len(text),
                "chunk_id": 0,
                "total_chunks": 1
            }
        )
        documents.append(doc)
    
    return documents


if __name__ == "__main__":
    # Test the document loader
    loader = DocumentLoader()
    
    print("📚 Document Loader Test")
    print("=" * 40)
    print(f"Supported formats: {loader.get_supported_formats()}")
    print(f"Chunk size: {loader.chunk_size}")
    print(f"Chunk overlap: {loader.chunk_overlap}")
    
    # Test with sample documents
    sample_docs = create_sample_documents()
    print(f"\nCreated {len(sample_docs)} sample documents")
    
    for i, doc in enumerate(sample_docs[:2]):  # Show first 2
        print(f"\nDocument {i+1}:")
        print(f"Content: {doc.page_content[:100]}...")
        print(f"Metadata: {doc.metadata}")
